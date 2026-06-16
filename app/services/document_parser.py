"""
Document parsing service for multiple file formats.
Supports PDF, TXT, DOC, DOCX, XLS, XLSX with async batch processing.
"""

import hashlib
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import chardet
import fitz  # PyMuPDF
import openpyxl
import PyPDF2
from docx import Document
from loguru import logger

from app.config import get_settings
from app.services.text_splitter import MarkdownChunker
from app.utils.async_helpers import AsyncBatchProcessor


class DocumentChunk:
    """Represents a chunk of text from a document."""

    def __init__(
        self,
        content: str,
        metadata: Dict[str, Any],
        chunk_index: int = 0,
        text_start: int = 0,
        text_end: int = 0
    ):
        self.content = content
        self.metadata = metadata
        self.chunk_index = chunk_index
        self.text_start = text_start
        self.text_end = text_end


class DocumentParser:
    """Parse documents of various formats into text chunks."""

    def __init__(self):
        self.settings = get_settings()
        self.chunk_size = self.settings.CHUNK_SIZE
        self.chunk_overlap = self.settings.CHUNK_OVERLAP
        self.batch_processor = AsyncBatchProcessor(
            batch_size=self.settings.BATCH_SIZE,
            max_workers=self.settings.MAX_WORKERS
        )

    async def parse_files(self, file_paths: List[str]) -> List[DocumentChunk]:
        """
        Parse multiple files and return document chunks.

        Args:
            file_paths: List of file paths to parse

        Returns:
            List of DocumentChunk objects
        """
        logger.info(f"Starting to parse {len(file_paths)} files")

        # Parse files in batches using thread pool
        all_chunks = await self.batch_processor.process_in_batches(
            items=file_paths,
            process_func=self._parse_single_file_sync,
            use_threads=True
        )

        # Flatten the list of chunks
        flattened_chunks = []
        for file_chunks in all_chunks:
            if file_chunks:  # Skip None or empty results
                flattened_chunks.extend(file_chunks)

        logger.info(f"Parsed {len(flattened_chunks)} total chunks from {len(file_paths)} files")
        return flattened_chunks

    def _parse_single_file_sync(self, file_path: str) -> List[DocumentChunk]:
        """
        Parse a single file (synchronous function for thread pool).

        Args:
            file_path: Path to the file

        Returns:
            List of DocumentChunk objects
        """
        try:
            file_path_obj = Path(file_path)

            if not file_path_obj.exists():
                logger.error(f"File not found: {file_path_obj.name}")
                return []

            # Check file size
            file_size_mb = file_path_obj.stat().st_size / (1024 * 1024)
            if file_size_mb > self.settings.MAX_FILE_SIZE_MB:
                logger.error(f"File too large ({file_size_mb:.2f}MB): {file_path_obj.name}")
                return []

            extension = file_path_obj.suffix.lower()
            logger.debug(f"Parsing {extension} file: {file_path_obj.name}")

            # Route to appropriate parser
            if extension == '.pdf':
                text = self._parse_pdf(file_path)
            elif extension == '.txt':
                text = self._parse_txt(file_path)
            elif extension in ['.doc', '.docx']:
                text = self._parse_docx(file_path)
            elif extension in ['.xls', '.xlsx']:
                text = self._parse_excel(file_path)
            else:
                logger.warning(f"Unsupported file format: {extension}")
                return []

            if not text or not text.strip():
                logger.warning(f"No text extracted from: {file_path_obj.name}")
                return []

            # Create chunks from text
            chunks = self._create_chunks(text, file_path)
            logger.info(f"Created {len(chunks)} chunks from {file_path_obj.name}")

            return chunks

        except Exception as e:
            logger.error(f"Error parsing {file_path_obj.name if 'file_path_obj' in locals() else 'file'}: {str(e)}")
            return []

    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF file using PyMuPDF (faster and more accurate)."""
        try:
            text_parts = []
            doc = fitz.open(file_path)

            for page_num, page in enumerate(doc, 1):
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"[Page {page_num}]\n{text}")

            doc.close()
            return "\n\n".join(text_parts)

        except Exception as e:
            logger.warning(f"PyMuPDF failed, trying PyPDF2: {str(e)}")
            # Fallback to PyPDF2
            try:
                text_parts = []
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(reader.pages, 1):
                        text = page.extract_text()
                        if text.strip():
                            text_parts.append(f"[Page {page_num}]\n{text}")
                return "\n\n".join(text_parts)
            except Exception as e2:
                logger.error(f"Both PDF parsers failed: {str(e2)}")
                return ""

    def _parse_txt(self, file_path: str) -> str:
        """Parse text file with encoding detection."""
        try:
            # Detect encoding
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'utf-8'

            # Read with detected encoding
            with open(file_path, 'r', encoding=encoding, errors='ignore') as file:
                return file.read()

        except Exception as e:
            logger.error(f"Error parsing text file: {str(e)}")
            return ""

    def _parse_docx(self, file_path: str) -> str:
        """Parse DOCX file."""
        try:
            doc = Document(file_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error parsing DOCX file: {str(e)}")
            return ""

    def _parse_excel(self, file_path: str) -> str:
        """Parse Excel file (XLS/XLSX)."""
        try:
            workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            text_parts = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"[Sheet: {sheet_name}]")

                for row in sheet.iter_rows(values_only=True):
                    row_text = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                    if row_text.strip():
                        text_parts.append(row_text)

            workbook.close()
            return "\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error parsing Excel file: {str(e)}")
            return ""

    def _create_chunks(self, text: str, source_file: str) -> List[DocumentChunk]:
        """
        Split text into structure-aware chunks prioritizing paragraph, heading,
        and natural boundary cuts before falling back to hard character cuts.

        Args:
            text: Full text to chunk
            source_file: Source file path

        Returns:
            List of DocumentChunk objects
        """
        file_path_obj = Path(source_file)
        document_id = self._generate_document_id(file_path_obj)

        # Use MarkdownChunker for structure-aware splitting
        if source_file.endswith('.md'):
            markdown_chunker = MarkdownChunker(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap
            )
            chunk_results = markdown_chunker.chunk(text)

            chunks_list = []
            for result in chunk_results:
                chunks_list.append(
                    DocumentChunk(
                        content=result.content,
                        metadata=self._build_metadata(
                            file_path_obj=file_path_obj,
                            document_id=document_id,
                            chunk_index=result.metadata.get("position", 0),
                            total_chunks=0,
                            text_start=0,
                            text_end=len(result.content),
                            chunk_type=result.metadata.get("chunk_type", "paragraph"),
                            heading_path=result.metadata.get("heading_path", []),
                        ),
                        chunk_index=result.metadata.get("position", 0),
                        text_start=0,
                        text_end=len(result.content),
                    )
                )

            total_chunks = len(chunks_list)
            for i, chunk in enumerate(chunks_list):
                chunk.metadata["total_chunks"] = total_chunks

            return chunks_list

        # Fallback to original behavior for non-Markdown files
        structural_segments = self._split_by_structure(text)

        chunks_list = []
        chunk_index = 0

        for segment, orig_start, orig_end in structural_segments:
            segment_len = orig_end - orig_start

            if segment_len <= self.chunk_size:
                if segment.strip():
                    chunks_list.append(
                        DocumentChunk(
                            content=segment.strip(),
                            metadata=self._build_metadata(
                                file_path_obj=file_path_obj,
                                document_id=document_id,
                                chunk_index=chunk_index,
                                total_chunks=0,
                                text_start=orig_start,
                                text_end=orig_end,
                                chunk_type="paragraph",
                                heading_path=[],
                            ),
                            chunk_index=chunk_index,
                            text_start=orig_start,
                            text_end=orig_end,
                        )
                    )
                    chunk_index += 1
            else:
                sub_chunks = self._split_large_segment(
                    segment,
                    file_path_obj,
                    document_id,
                    chunk_index,
                    orig_start,
                    orig_end,
                )
                chunks_list.extend(sub_chunks)
                chunk_index += len(sub_chunks)

        total_chunks = len(chunks_list)
        for i, chunk in enumerate(chunks_list):
            chunk.metadata["total_chunks"] = total_chunks

        return chunks_list

    def _split_by_structure(self, text: str) -> List[Tuple[str, int, int]]:
        """
        Split text by structural boundaries: paragraphs, then lines.
        Returns list of (segment, start_offset, end_offset) tuples preserving original positions.
        """
        segments = []

        double_newline_pattern = r'\n\n+'
        matches = list(re.finditer(double_newline_pattern, text))

        if not matches:
            if len(text) <= self.chunk_size * 1.5:
                if text.strip():
                    segments.append((text, 0, len(text)))
            else:
                single_newline_pattern = r'\n'
                sub_matches = list(re.finditer(single_newline_pattern, text))
                prev_end = 0
                for m in sub_matches:
                    segment = text[prev_end:m.start()]
                    if segment.strip():
                        segments.append((segment, prev_end, m.start()))
                    prev_end = m.end()
                remaining = text[prev_end:]
                if remaining.strip():
                    segments.append((remaining, prev_end, len(text)))
            return segments

        prev_pos = 0
        for m in matches:
            segment = text[prev_pos:m.start()]
            if segment.strip():
                segments.append((segment, prev_pos, m.start()))
            prev_pos = m.end()

        remaining = text[prev_pos:]
        if remaining.strip():
            segments.append((remaining, prev_pos, len(text)))

        processed = []
        for segment, start, end in segments:
            if len(segment) <= self.chunk_size * 1.5:
                processed.append((segment, start, end))
            else:
                single_newline_pattern = r'\n'
                sub_matches = list(re.finditer(single_newline_pattern, segment))
                if not sub_matches:
                    processed.append((segment, start, end))
                else:
                    prev_sub = 0
                    for sm in sub_matches:
                        sub_seg = segment[prev_sub:sm.start()]
                        if sub_seg.strip():
                            processed.append((sub_seg, start + prev_sub, start + sm.start()))
                        prev_sub = sm.end()
                    remaining_sub = segment[prev_sub:]
                    if remaining_sub.strip():
                        processed.append((remaining_sub, start + prev_sub, end))

        return processed

    def _split_large_segment(
        self,
        segment: str,
        file_path_obj: Path,
        document_id: str,
        start_chunk_index: int,
        orig_start: int,
        orig_end: int,
    ) -> List[DocumentChunk]:
        """
        Split a large segment that exceeds chunk_size using
        sentence and word boundaries.
        """
        chunks_list = []
        char_position = orig_start
        chunk_index = start_chunk_index

        while char_position < orig_end:
            remaining_len = orig_end - char_position
            if remaining_len <= self.chunk_size:
                chunks_list.append(
                    DocumentChunk(
                        content=segment[char_position - orig_start:].strip(),
                        metadata=self._build_metadata(
                            file_path_obj=file_path_obj,
                            document_id=document_id,
                            chunk_index=chunk_index,
                            total_chunks=0,
                            text_start=char_position,
                            text_end=orig_end,
                        ),
                        chunk_index=chunk_index,
                        text_start=char_position,
                        text_end=orig_end,
                    )
                )
                break

            chunk_text = segment[char_position - orig_start:char_position - orig_start + self.chunk_size]

            break_point = self._find_best_break_point(chunk_text)

            if break_point > self.chunk_size * 0.4:
                actual_end = char_position + break_point + 1
                chunk_text = chunk_text[:break_point + 1]
            else:
                actual_end = char_position + self.chunk_size
                chunk_text = chunk_text.strip()

            if chunk_text:
                chunks_list.append(
                    DocumentChunk(
                        content=chunk_text.strip(),
                        metadata=self._build_metadata(
                            file_path_obj=file_path_obj,
                            document_id=document_id,
                            chunk_index=chunk_index,
                            total_chunks=0,
                            text_start=char_position,
                            text_end=actual_end,
                        ),
                        chunk_index=chunk_index,
                        text_start=char_position,
                        text_end=actual_end,
                    )
                )
                chunk_index += 1

            char_position = actual_end - self.chunk_overlap

        return chunks_list

    def _find_best_break_point(self, chunk_text: str) -> int:
        """
        Find the best break point within chunk_text, preferring
        sentence boundaries > paragraph boundaries > word boundaries.
        """
        sentence_ends = list(re.finditer(r'[.!?]+\s+', chunk_text))
        if sentence_ends:
            last_sentence = sentence_ends[-1]
            if last_sentence.start() > self.chunk_size * 0.5:
                return last_sentence.start()

        newline_pos = chunk_text.rfind('\n')
        if newline_pos > self.chunk_size * 0.5:
            return newline_pos

        space_pos = chunk_text.rfind(' ')
        if space_pos > self.chunk_size * 0.5:
            return space_pos

        comma_pos = chunk_text.rfind(',')
        if comma_pos > self.chunk_size * 0.7:
            return comma_pos

        return -1

    def _generate_document_id(self, file_path_obj: Path) -> str:
        """Generate a stable document ID from file path and name."""
        try:
            mtime = file_path_obj.stat().st_mtime
        except OSError:
            mtime = 0
        id_input = f"{file_path_obj.name}:{mtime}"
        return hashlib.sha256(id_input.encode()).hexdigest()[:16]

    def _build_metadata(
        self,
        file_path_obj: Path,
        document_id: str,
        chunk_index: int,
        total_chunks: int,
        text_start: int,
        text_end: int,
        chunk_type: str = "paragraph",
        heading_path: Optional[List[str]] = None,
    ) -> Dict[str, Any]:
        """Build enriched metadata for a chunk."""
        return {
            'document_id': document_id,
            'source_file': file_path_obj.name,
            'file_path': file_path_obj.name,
            'file_type': file_path_obj.suffix.lower(),
            'chunk_index': chunk_index,
            'total_chunks': total_chunks,
            'text_start': text_start,
            'text_end': text_end,
            'chunk_char_length': text_end - text_start,
            'chunk_type': chunk_type,
            'heading_path': heading_path or [],
        }
